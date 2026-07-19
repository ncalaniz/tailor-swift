# export.py — turns tailored resume text into downloadable Word (.docx) and PDF files.
import io
import docx                 # installed as "python-docx"
from fpdf import FPDF       # installed as "fpdf2"
import store
import re
import json

def _contact():
    """Pull the header (name + contact line) from settings."""
    name = store.get_setting("name", "")
    bits = [store.get_setting("email", ""), store.get_setting("phone", ""),
            store.get_setting("location", ""), store.get_setting("linkedin", "")]
    return name, [b for b in bits if b]

def _is_rule(line):
    """True for leftover '---' style separator lines we want to drop."""
    return len(line) >= 3 and set(line) <= {"-", "*", "_"}

def _extract_headline(text):
    """Pull the '## Headline' line out of the tailored text; return (headline, rest)."""
    lines = text.split("\n")
    headline = ""
    out = []
    i = 0
    while i < len(lines):
        if lines[i].strip() == "## Headline":
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines):
                headline = lines[i].strip()
                i += 1
            continue
        out.append(lines[i])
        i += 1
    return headline, "\n".join(out)

def _ascii(s):
    """Swap characters the PDF font can't handle for plain equivalents."""
    for bad, good in [("—", "-"), ("–", "-"), ("’", "'"), ("‘", "'"),
                      ("“", '"'), ("”", '"'), ("•", "-"), ("…", "...")]:
        s = s.replace(bad, good)
    return s.encode("latin-1", "replace").decode("latin-1")

def _norm(s):
    """Lowercase + collapse whitespace, for tolerant matching."""
    return re.sub(r"\s+", " ", (s or "").lower()).strip()

def _dates_for_heading(heading):
    """Find the stored job inside this heading and return a formatted date range (or '')."""
    h = _norm(heading)
    # longest role first, so 'Sr. Manager...' wins over 'Manager...'
    for j in sorted(store.list_jobs(), key=lambda x: -len(x["role"] or "")):
        emp, role = _norm(j["employer"]), _norm(j["role"])
        if emp and emp in h and (not role or role in h):
            start = (j["start_date"] or "").strip()
            end = (j["end_date"] or "").strip()
            if start and end:
                return f"{start} \u2013 {end}"
            if start:
                return f"{start} \u2013 Present"
            return end
    return ""

def build_docx(tailored_text):
    """Turn the tailored text into a Word doc, returned as bytes."""
    headline, tailored_text = _extract_headline(tailored_text)
    doc = docx.Document()
    name, contact_parts = _contact()
    if name:
        doc.add_heading(name, level=0)
    for part in contact_parts:
        doc.add_paragraph(part)
    if headline:
        p = doc.add_paragraph()
        run = p.add_run(headline)
        run.bold = True

    for raw in tailored_text.split("\n"):
        line = raw.strip()
        if not line or _is_rule(line):
            continue
        job_id = _job_tag(line)
        if job_id is not None:
            title, dates, loc = _job_heading_parts(job_id)
            if title:
                doc.add_heading(title + (f"   ({dates})" if dates else ""), level=2)
                if loc:
                    doc.add_paragraph(loc)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    skills = _skills()
    if skills:
        doc.add_heading("Technologies", level=1)
        doc.add_paragraph(", ".join(skills))

    edu = _education()
    if edu:
        doc.add_heading("Education", level=1)
        for e in edu:
            school = e.get("school", "")
            loc = e.get("location", "")
            degree = e.get("degree", "")
            year = e.get("year", "")
            doc.add_heading(school + (f" ({loc})" if loc else ""), level=2)
            line2 = degree + (f", {year}" if year else "")
            if line2.strip():
                doc.add_paragraph(line2)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def _job_heading_parts(job_id):
    """Return (title_text, dates_text, location_text) for a job id."""
    for j in store.list_jobs():
        if j["id"] == job_id:
            title = f"{j['employer']} — {j['role']}"
            start = (j["start_date"] or "").strip()
            end = (j["end_date"] or "").strip()
            dates = f"{start or '?'} \u2013 {end or 'Present'}" if (start or end) else ""
            loc = (j["location"] or "").strip()
            return title, dates, loc
    return "", "", ""

def _job_tag(line):
    """If the line is a [[JOB:id]] tag, return the id (int); else None."""
    m = re.match(r"^\[\[JOB:(\d+)\]\]\s*$", line.strip())
    return int(m.group(1)) if m else None

def _education():
    """Return the stored education list (or empty)."""
    try:
        return json.loads(store.get_setting("education", "[]") or "[]")
    except Exception:
        return []

def _skills():
    """Return the stored skills list (or empty)."""
    try:
        return json.loads(store.get_setting("skills", "[]") or "[]")
    except Exception:
        return []

def _section_header(pdf, title):
    """Draw a section title with a horizontal rule under it."""
    pdf.set_font("Helvetica", "B", 15)
    pdf.multi_cell(0, 8, _ascii(title), new_x="LMARGIN", new_y="NEXT")
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(2)

_MONTHS = {"jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
           "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12}

def _year_of(text):
    """Pull a 4-digit year out of a date string like 'Apr 2023' (0 if none)."""
    m = re.search(r"(\d{4})", text or "")
    return int(m.group(1)) if m else 0

def _month_of(text):
    """Pull a month number out of a date string like 'Apr 2023' (6 = mid-year, if none)."""
    t = (text or "").strip().lower()[:3]
    return _MONTHS.get(t, 6)

def build_pdf(tailored_text):
    """Turn the tailored text into a PDF, returned as bytes."""
    headline, tailored_text = _extract_headline(tailored_text)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    name, contact_parts = _contact()
    if name:
        pdf.set_font("Helvetica", "B", 20)
        pdf.multi_cell(0, 9, _ascii(name), new_x="LMARGIN", new_y="NEXT")
    if contact_parts:
        pdf.set_font("Helvetica", "", 10)
        for part in contact_parts:
            pdf.cell(0, 5, _ascii(part), new_x="LMARGIN", new_y="NEXT")
    if headline:
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(0, 6, _ascii(headline), new_x="LMARGIN", new_y="NEXT", align="L")
    pdf.ln(4)

    # --- pass 1: split the text into preamble (summary) and per-job sections ---
    preamble = []          # lines before the first job tag
    sections = []          # list of (job_id, [bullet lines])
    current = None         # the bullet list of the job we're inside
    for raw in tailored_text.split("\n"):
        line = raw.strip()
        if not line or _is_rule(line):
            continue
        job_id = _job_tag(line)
        if job_id is not None:
            current = []
            sections.append((job_id, current))
            continue
        if current is None:
            preamble.append(line)
        else:
            current.append(line)

    # --- sort sections most-recent-first by the job's dates ---
    def _sort_key(sec):
        job_id, _ = sec
        for j in store.list_jobs():
            if j["id"] == job_id:
                end = (j["end_date"] or "").strip().lower()
                start = (j["start_date"] or "").strip()
                if not end or end == "present":
                    return (0, "")            # current job: always first
                return (1, (_year_of(end) * -1, _month_of(end) * -1))  # newest first, month breaks ties
        return (2, 0)
    sections.sort(key=_sort_key)

    # --- pass 2: draw ---
    def _draw_line(line):
        if line.startswith("### "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 7, _ascii(line[4:]), new_x="LMARGIN", new_y="NEXT")
        elif line.startswith("## "):
            _section_header(pdf, line[3:])
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 5.5, chr(149) + " " + _ascii(line[2:]), new_x="LMARGIN", new_y="NEXT", align="L")
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 5.5, _ascii(line), new_x="LMARGIN", new_y="NEXT", align="L")
        pdf.ln(1.5)

    for line in preamble:
        _draw_line(line)

    drew_experience = False
    for job_id, bullets in sections:
        title, dates, loc = _job_heading_parts(job_id)
        if not title:
            continue
        if not drew_experience:
            pdf.ln(2)
            _section_header(pdf, "Experience")
            drew_experience = True
        pdf.set_font("Helvetica", "", 11)
        dates_w = pdf.get_string_width(_ascii(dates)) + 4
        pdf.set_font("Helvetica", "B", 13)
        title_txt = _ascii(title)
        usable = pdf.w - pdf.l_margin - pdf.r_margin
        max_title_w = usable - dates_w
        if pdf.get_string_width(title_txt) + 2 <= max_title_w:
            # fits on one line: title left, dates right
            title_w = pdf.get_string_width(title_txt) + 2
            pdf.cell(title_w, 7, title_txt)
            pdf.set_font("Helvetica", "", 11)
            pdf.cell(0, 7, _ascii(dates), align="R", new_x="LMARGIN", new_y="NEXT")
        else:
            # too long to share a line: let the title wrap, put dates underneath
            pdf.multi_cell(0, 7, title_txt, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 11)
            pdf.cell(0, 6, _ascii(dates), align="R", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 11)
        if loc:
            pdf.set_font("Helvetica", "I", 10)
            pdf.cell(0, 5, _ascii(loc), align="R", new_x="LMARGIN", new_y="NEXT")
        for line in bullets:
            _draw_line(line)

    # --- skills ---
    skills = _skills()
    if skills:
        pdf.ln(3)
        _section_header(pdf, "Technologies")
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 5.5, _ascii(", ".join(skills)), new_x="LMARGIN", new_y="NEXT", align="L")

    # --- education (unchanged) ---
    edu = _education()
    if edu:
        pdf.ln(3)
        _section_header(pdf, "Education")
        for e in edu:
            school = e.get("school", "")
            degree = e.get("degree", "")
            loc = e.get("location", "")
            year = e.get("year", "")
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 6, _ascii(school + (f"   ({loc})" if loc else "")), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 11)
            line2 = degree + (f"   {year}" if year else "")
            if line2.strip():
                pdf.multi_cell(0, 6, _ascii(line2), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

    return bytes(pdf.output())
